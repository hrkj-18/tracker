# encoding: utf-8
import os
from dashboard.utils import Details
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_doc(work_item):
    try:
        id = work_item.id
        title = work_item.title
        description = work_item.description
        owner = work_item.owner
        org = Details.ORGANIZATION
        company = Details.COMPANY
        business_unit = Details.BUSINESS_UNIT

        document = Document()

        # Add Header
        section = document.sections[0]
        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = "Work Package"
        header_para.text += f"\t\t{id}_{title}\n\t\t{company}"

        # Add Title
        doc_title = document.add_heading(title, 0)
        doc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        document.add_heading(f'Section 1 - Request from {org}', level=1)

        p = document.add_paragraph()
        # run = p.add_run()
        # run.add_break()

        table_1 = document.add_table(rows=5, cols=2)
        table_1.style = 'TableGrid'
        cols_0 = table_1.columns[0].cells
        cols_0[0].text = 'Unique ID'
        cols_0[1].text = 'Title of Incident'
        cols_0[2].text = 'Business Unit'
        cols_0[3].text = 'Request Raised by'
        cols_0[4].text = 'Date Request Raised'
        cols_1 = table_1.columns[1].cells
        cols_1[0].text = str(id)
        cols_1[1].text = title
        cols_1[2].text = business_unit
        cols_1[3].text = owner
        cols_1[4].text = work_item.created.strftime('%d.%m.%Y %H:%M:%S')

        document.add_paragraph()

        table_2 = document.add_table(rows=2, cols=1)
        table_2.style = 'TableGrid'
        table_2.columns[0].cells[0].text = 'Detailed Description'
        table_2.columns[0].cells[1].text = description

        # Adding a page break
        document.add_page_break()

        document.add_heading('Section 2 - Impact Assesment', level=1)
        document.add_paragraph("The proposed solution and estimate for this change request is detailed below.")
        table_3 = document.add_table(rows=3, cols=2)
        table_3.style = 'TableGrid'
        cols_0 = table_3.columns[0].cells
        cols_0[0].text = 'Completed By'
        cols_0[1].text = 'Date'
        cols_0[2].text = 'Version'

        document.add_heading('2.1 Proposed Solution', level=2)

        # Adding a page break
        document.add_page_break()

        document.add_heading('2.2 Assumptions', level=2)
        document.add_paragraph(
            f'Deployment is to be done by {company}', style='List Number'
        )
        document.add_paragraph()

        document.add_heading('2.3 Work Products', level=2)
        document.add_paragraph(
            f'In order to deliver the changes outlined in Section 2 of the Impact Assessment, the following work products will be produced by {company}:'
        )
        document.add_paragraph(
            'Impact Analysis Document', style='List Bullet'
        )
        document.add_paragraph()

        document.add_heading('2.4 GDPR Checkpoints', level=2)

        table_4 = document.add_table(rows=5, cols=3)
        table_4.style = 'TableGrid'
        col_1 = table_4.columns[0].cells
        col_1[0].text = 'Date'
        col_1[1].text = '1. Do the requirements of this change impact data under GDPR?'
        col_1[2].text = '2. Is this data object(s) available in UAT for user testing sign off?'
        col_1[3].text = '3. If the data object(s) is not avaiable, how will this change be tested in UAT by user?'
        col_1[4].text = '4. Please include additional effort on this change where daa preparation is required for testing sign off'
        col_2 = table_4.columns[1].cells
        col_2[0].text = 'Yes/No?'
        col_3 = table_4.columns[2].cells
        col_3[0].text = 'If Yes - Please mention detail of impact'

        # Adding a page break
        document.add_page_break()

        document.add_heading('Section 3 - Reviewers/Approvers', level=1)

        table_5 = document.add_table(rows=4, cols=6)
        table_5.style = 'TableGrid'
        header_cells = table_5.rows[0].cells
        header_cells[0].text = 'Name'
        header_cells[1].text = 'Approver/Reviewer'
        header_cells[2].text = 'Approved/Reviewed'
        header_cells[3].text = 'Comments'
        header_cells[4].text = 'Date'
        header_cells[5].text = 'Doc Version'

        path = os.path.join(os.getcwd(), 'docs\\IA')
        if not os.path.exists(path):
            os.makedirs(path)

        title = '-'.join(title.split(' '))
        file_name = f"IA-{id}-{title}.docx"
        file_path = os.path.join(path, file_name)

        document.save(file_path)

        return {'message': "IA Document Created",'file_path': file_path, 'file_name': file_name}

    except Exception as e:
        print(e)
        return {'message': "Cannot Create Document",'file_path': None, 'file_name': None}
